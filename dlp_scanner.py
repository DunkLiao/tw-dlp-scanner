# -*- coding: utf-8 -*-
"""
DLP Scanner - 敏感資料掃描工具

功能：
1. 使用 tkinter + ttk 建立圖形化介面
2. 可選擇指定資料夾掃描
3. 可掃描子資料夾
4. 支援 txt/csv/log/md/json/xml/html/docx/xlsx/pdf
5. 偵測：
   - 組合式個資
   - 加密檔案
   - 台灣地址
   - 中文姓名
   - 台灣身分證字號
   - 護照號碼
   - 信用卡號，含 Luhn 檢核
   - 銀行存摺根號與關鍵字
   - 公文密等關鍵字組合
   - 電子郵件
   - 居留證號碼
6. 可匯出 Excel 掃描報表
"""

import os
import re
import sys
import tempfile
import zipfile
import threading
import queue
import subprocess
from pathlib import Path
from datetime import datetime
from tkinter import Tk, StringVar, BooleanVar, filedialog, messagebox, END
from tkinter import ttk

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook, Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
except ImportError:
    load_workbook = None
    Workbook = None
    Font = None
    PatternFill = None
    Alignment = None
    get_column_letter = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    import olefile
except ImportError:
    olefile = None

APP_NAME = "DLP 敏感資料掃描工具"
APP_VERSION = "1.0.0"

SUPPORTED_TEXT_EXTS = {
    ".txt", ".csv", ".log", ".md", ".json", ".xml", ".html", ".htm"
}
SUPPORTED_DOC_EXTS = {".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pdf"}
SUPPORTED_SCAN_EXTS = SUPPORTED_TEXT_EXTS | SUPPORTED_DOC_EXTS
ENCRYPT_TARGET_EXTS = {".docx", ".xlsx", ".pptx", ".pdf", ".zip", ".rar", ".7z"}
ENCRYPTED_FILE_RULE = "2. 加密檔案"
MAX_ZIP_SCAN_DEPTH = 5
MAX_ZIP_ENTRY_BYTES = 50 * 1024 * 1024
MAX_ZIP_TOTAL_BYTES = 500 * 1024 * 1024

ACTION_BY_RISK = {
    "低": "僅記錄",
    "中": "僅記錄",
    "高": "記錄並阻擋",
    "未知": "待確認",
}

POLICY_THRESHOLDS = {
    "3. 台灣地址": {"低": 40, "中": 60, "高": 80},
    "4. 中文姓名": {"低": 40, "中": 80, "高": 100},
    "5. 台灣身分證字號": {"中": 1, "高": 2},
    "6. 護照號碼": {"中": 1, "高": 2},
    "7. 信用卡號": {"中": 1, "高": 2},
    "10. 電子郵件": {"低": 40, "中": 60, "高": 80},
    "11. 居留證號碼": {"中": 1, "高": 2},
}

TAIWAN_CITY_COUNTY = (
    "台北市|臺北市|新北市|桃園市|台中市|臺中市|台南市|臺南市|高雄市|"
    "基隆市|新竹市|嘉義市|新竹縣|苗栗縣|彰化縣|南投縣|雲林縣|"
    "嘉義縣|屏東縣|宜蘭縣|花蓮縣|台東縣|臺東縣|澎湖縣|金門縣|連江縣"
)

SINGLE_PATTERNS = {
    "3. 台灣地址": re.compile(
        rf"({TAIWAN_CITY_COUNTY})"
        r"[\u4e00-\u9fffA-Za-z0-9]{0,30}"
        r"(路|街|大道|巷|弄|段)"
        r"[\u4e00-\u9fffA-Za-z0-9\-之號樓室F\s]{0,30}"
    ),
    "4. 中文姓名": re.compile(
        r"(?<![\u4e00-\u9fff])"
        r"[\u4e00-\u9fff]{2,4}"
        r"(?=(先生|小姐|女士|經理|副理|襄理|課長|科長|處長|主任|同仁|客戶|君|收|敬啟))"
    ),
    "5. 台灣身分證字號": re.compile(r"\b[A-Z][12]\d{8}\b"),
    "6. 護照號碼": re.compile(r"\b\d{9}\b"),
    "10. 電子郵件": re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"),
    "11. 居留證號碼": re.compile(r"\b[A-Z][A-D89]\d{8}\b"),
    "手機號碼": re.compile(r"\b09\d{2}[- ]?\d{3}[- ]?\d{3}\b"),
    "電話號碼": re.compile(r"\b0[2-8][- ]?\d{3,4}[- ]?\d{4}\b"),
    "Taiwan ID 關鍵字": re.compile(r"Taiwan\s*ID", re.IGNORECASE),
}

CREDIT_CARD_PATTERN = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
BANK_ACCOUNT_PATTERN = re.compile(r"(?<!\d)\d{12,14}(?!\d)")
BANK_KEYWORD_PATTERN = re.compile(r"客戶往來明細查詢")

COMBO_RULES = [
    {
        "name": "9. 公文密等關鍵字組合",
        "all_of": [
            re.compile(r"發文字號"),
            re.compile(r"(密等及解密條件|保密期限)"),
            re.compile(r"(極機密|機密|密件|密\s)"),
        ],
    },
]

COMBO_PI_PAIRS = [
    ("4. 中文姓名", "手機號碼"),
    ("4. 中文姓名", "3. 台灣地址"),
    ("4. 中文姓名", "電話號碼"),
    ("4. 中文姓名", "Taiwan ID 關鍵字"),
    ("3. 台灣地址", "Taiwan ID 關鍵字"),
]


def normalize_sample(value):
    """將 regex findall 回傳值整理為可顯示字串。"""
    if isinstance(value, tuple):
        return "".join(str(x) for x in value if x)
    return str(value)


def mask_sensitive(sample):
    """報表與畫面顯示時遮罩部分敏感內容，降低二次外洩風險。"""
    text = str(sample)

    if "@" in text:
        return re.sub(
            r"([A-Za-z0-9._%+\-]{2})[A-Za-z0-9._%+\-]*(@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})",
            r"\1***\2",
            text,
        )

    if re.fullmatch(r"[A-Z][A-D8912]\d{8}", text):
        return text[:3] + "*****" + text[-2:]

    digits = re.sub(r"\D", "", text)
    if len(digits) >= 8:
        return re.sub(r"\d(?=\d{4})", "*", text)

    return text


def luhn_check(card_no):
    """信用卡 Luhn 演算法檢核。"""
    card_no = re.sub(r"\D", "", str(card_no))
    if not card_no.isdigit() or not 13 <= len(card_no) <= 19:
        return False

    checksum = 0
    for index, char in enumerate(card_no[::-1]):
        digit = int(char)
        if index % 2 == 1:
            digit *= 2
            if digit > 9:
                digit -= 9
        checksum += digit
    return checksum % 10 == 0


def find_credit_cards(text):
    """先以 regex 找出疑似信用卡號，再使用 Luhn 檢核。"""
    results = []
    for match in CREDIT_CARD_PATTERN.finditer(text):
        raw = match.group()
        if luhn_check(raw):
            results.append(raw)
    return results


def risk_from_thresholds(count, thresholds):
    for risk in ("高", "中", "低"):
        threshold = thresholds.get(risk)
        if threshold is not None and count >= threshold:
            return risk
    return None


def make_policy_hit(risk, count, samples):
    return {
        "risk": risk,
        "action": ACTION_BY_RISK.get(risk, "待確認"),
        "hit_count": count,
        "samples": samples[:5],
    }


def collect_pattern_hits(pattern, text):
    matches = [match.group(0) for match in pattern.finditer(text)]
    return {
        "count": len(matches),
        "samples": matches[:5],
    }


def collect_text_detections(text):
    detected = {}

    for name, pattern in SINGLE_PATTERNS.items():
        hit = collect_pattern_hits(pattern, text)
        if hit["count"]:
            detected[name] = hit

    cards = find_credit_cards(text)
    if cards:
        detected["7. 信用卡號"] = {
            "count": len(cards),
            "samples": cards[:5],
        }

    return detected


def has_combo_personal_data(detected):
    for left, right in COMBO_PI_PAIRS:
        left_hit = detected.get(left)
        right_hit = detected.get(right)
        if left_hit and right_hit and left_hit["count"] >= 2 and right_hit["count"] >= 2:
            return True
    return False


def detect_or_rule_hits(text):
    detected = collect_text_detections(text)
    hit_names = []

    if has_combo_personal_data(detected):
        hit_names.append("1. 組合式個資")

    for name in [
        "3. 台灣地址",
        "4. 中文姓名",
        "5. 台灣身分證字號",
        "6. 護照號碼",
        "7. 信用卡號",
    ]:
        if name in detected:
            hit_names.append(name)

    if BANK_ACCOUNT_PATTERN.search(text):
        hit_names.append("8. 銀行存款帳號與關鍵字")

    for rule in COMBO_RULES:
        if all(pattern.search(text) for pattern in rule["all_of"]):
            hit_names.append(rule["name"])

    for name in ["10. 電子郵件", "11. 居留證號碼"]:
        if name in detected:
            hit_names.append(name)

    return hit_names


def collect_or_hit_samples(text, hit_names):
    detected = collect_text_detections(text)
    samples_by_hit = {}

    for hit_name in hit_names:
        hit = detected.get(hit_name)
        if hit:
            samples_by_hit[hit_name] = hit["samples"]

    bank_hit_names = [hit_name for hit_name in hit_names if hit_name.startswith("8. ")]
    if bank_hit_names:
        accounts = [match.group(0) for match in BANK_ACCOUNT_PATTERN.finditer(text)]
        for hit_name in bank_hit_names:
            samples_by_hit[hit_name] = accounts[:5]

    return samples_by_hit


def read_text_file(path):
    try:
        data = Path(path).read_bytes()
    except Exception:
        return ""

    if not data:
        return ""

    bom_encodings = [
        (b"\xef\xbb\xbf", "utf-8-sig"),
        (b"\xff\xfe", "utf-16"),
        (b"\xfe\xff", "utf-16"),
    ]
    for bom, enc in bom_encodings:
        if data.startswith(bom):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                return ""

    encodings = ["utf-8", "cp950", "big5"]
    candidates = []
    for enc in encodings:
        try:
            text = data.decode(enc)
            candidates.append((score_decoded_text(text), text))
        except UnicodeDecodeError:
            continue
    if candidates:
        candidates.sort(key=lambda item: item[0], reverse=True)
        return candidates[0][1]

    for enc in ["utf-16", "latin-1"]:
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return ""


def score_decoded_text(text):
    if not text:
        return -1
    replacement_count = text.count("\ufffd")
    control_count = sum(1 for char in text if ord(char) < 32 and char not in "\r\n\t")
    cjk_count = sum(1 for char in text if "\u4e00" <= char <= "\u9fff")
    printable_count = sum(1 for char in text if char.isprintable() or char in "\r\n\t")
    return (cjk_count * 4) + printable_count - (replacement_count * 20) - (control_count * 10)


def read_docx(path):
    if Document is None:
        return ""
    try:
        doc = Document(path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception:
        return ""


def read_ole_text_from_bytes(data):
    text_parts = []

    for encoding in ("utf-16le", "utf-16be"):
        decoded = data.decode(encoding, errors="ignore")
        cleaned = "".join(char if (char.isprintable() or char in "\r\n\t") else "\n" for char in decoded)
        text_parts.extend(part.strip() for part in re.split(r"\s*\n+\s*", cleaned) if len(part.strip()) >= 4)

    for match in re.finditer(rb"[\x09\x0a\x0d\x20-\x7e]{4,}", data):
        text_parts.append(match.group().decode("latin-1", errors="ignore").strip())

    seen = set()
    unique_parts = []
    for part in text_parts:
        if part and part not in seen:
            seen.add(part)
            unique_parts.append(part)
    return "\n".join(unique_parts)


def read_ole_text(path):
    file_path = Path(path)
    parts = []

    if olefile is not None:
        try:
            if olefile.isOleFile(str(file_path)):
                with olefile.OleFileIO(str(file_path)) as ole:
                    for stream_name in ole.listdir(streams=True, storages=False):
                        try:
                            with ole.openstream(stream_name) as stream:
                                parts.append(read_ole_text_from_bytes(stream.read()))
                        except Exception:
                            continue
        except Exception:
            parts = []

    if not parts:
        try:
            parts.append(read_ole_text_from_bytes(file_path.read_bytes()))
        except Exception:
            return ""

    return "\n".join(part for part in parts if part)


def read_doc(path):
    return read_ole_text(path)


def read_ppt(path):
    return read_ole_text(path)


def read_xls(path):
    if xlrd is None:
        return ""
    try:
        workbook = xlrd.open_workbook(str(path), on_demand=True)
        rows = []
        try:
            for sheet in workbook.sheets():
                rows.append(f"[Sheet {sheet.name}]")
                for row_index in range(sheet.nrows):
                    values = sheet.row_values(row_index)
                    rows.append(" ".join("" if value is None else str(value) for value in values))
        finally:
            release = getattr(workbook, "release_resources", None)
            if release is not None:
                release()
        return "\n".join(rows)
    except Exception:
        return ""


def read_xlsx(path):
    if load_workbook is None:
        return ""
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        rows = []
        for sheet in workbook.worksheets:
            rows.append(f"[工作表] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                rows.append(" ".join("" if value is None else str(value) for value in row))
        workbook.close()
        return "\n".join(rows)
    except Exception:
        return ""


def read_pdf(path):
    if pdfplumber is None:
        return ""
    try:
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages_text.append(page.extract_text() or "")
        return "\n".join(pages_text)
    except Exception:
        return ""


def is_encrypted_file(path):
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            if pdfplumber is None:
                return False
            try:
                with pdfplumber.open(path) as pdf:
                    if len(pdf.pages) > 0:
                        _ = pdf.pages[0].extract_text()
                return False
            except Exception as err:
                message = str(err).lower()
                return "password" in message or "encrypted" in message

        if ext in {".docx", ".xlsx", ".pptx"}:
            try:
                with zipfile.ZipFile(path, "r"):
                    return False
            except zipfile.BadZipFile:
                return True
            except Exception:
                return False

        if ext == ".zip":
            try:
                with zipfile.ZipFile(path, "r") as zf:
                    return any(info.flag_bits & 0x1 for info in zf.infolist())
            except Exception:
                return False

        if ext in {".rar", ".7z"}:
            return False
    except Exception:
        return False
    return False


def extract_text(path):
    ext = path.suffix.lower()
    if ext in SUPPORTED_TEXT_EXTS:
        return read_text_file(path)
    if ext == ".doc":
        return read_doc(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".xls":
        return read_xls(path)
    if ext == ".xlsx":
        return read_xlsx(path)
    if ext == ".ppt":
        return read_ppt(path)
    if ext == ".pdf":
        return read_pdf(path)
    return ""


def scan_text(text):
    detected = collect_text_detections(text)
    hits = {}

    for name, thresholds in POLICY_THRESHOLDS.items():
        hit = detected.get(name)
        if not hit:
            continue
        risk = risk_from_thresholds(hit["count"], thresholds)
        if risk:
            hits[name] = make_policy_hit(risk, hit["count"], hit["samples"])

    if BANK_KEYWORD_PATTERN.search(text):
        accounts = [match.group(0) for match in BANK_ACCOUNT_PATTERN.finditer(text)]
        if accounts:
            hits["8. 銀行存款帳號與關鍵字"] = make_policy_hit(
                "高",
                len(accounts),
                accounts[:5],
            )

    for rule in COMBO_RULES:
        if all(pattern.search(text) for pattern in rule["all_of"]):
            hits[rule["name"]] = make_policy_hit("高", 1, ["關鍵字組合命中"])

    combo_candidates = []
    for left, right in COMBO_PI_PAIRS:
        left_hit = detected.get(left)
        right_hit = detected.get(right)
        if not left_hit or not right_hit:
            continue
        if left_hit["count"] < 2 or right_hit["count"] < 2:
            continue

        max_count = max(left_hit["count"], right_hit["count"])
        if max_count >= 5:
            risk = "高"
        elif max_count >= 3:
            risk = "中"
        else:
            risk = "低"

        combo_candidates.append({
            "risk": risk,
            "count": left_hit["count"] + right_hit["count"],
            "sample": f"{left}({left_hit['count']}) + {right}({right_hit['count']})",
        })

    if combo_candidates:
        combo_candidates.sort(key=lambda item: ("低", "中", "高").index(item["risk"]), reverse=True)
        best = combo_candidates[0]
        hits["1. 組合式個資"] = make_policy_hit(best["risk"], best["count"], [best["sample"]])

    return hits


def scan_file(path):
    hits = {}
    try:
        ext = path.suffix.lower()
        if ext in ENCRYPT_TARGET_EXTS and is_encrypted_file(path):
            hits[ENCRYPTED_FILE_RULE] = make_policy_hit("高", 1, [f"{ext} 可能為加密檔案"])

        if ext in SUPPORTED_SCAN_EXTS:
            text = extract_text(path)
            if text:
                hits.update(scan_text(text))
        return hits, ""
    except Exception as err:
        return hits, str(err)


def scan_file_with_or_hits(path):
    hits = {}
    or_hits = []
    try:
        ext = path.suffix.lower()
        if ext in ENCRYPT_TARGET_EXTS and is_encrypted_file(path):
            hits[ENCRYPTED_FILE_RULE] = make_policy_hit("高", 1, [f"{ext} 可能為加密檔案"])
            or_hits.append(ENCRYPTED_FILE_RULE)

        if ext in SUPPORTED_SCAN_EXTS:
            text = extract_text(path)
            if text:
                hits.update(scan_text(text))
                text_or_hits = detect_or_rule_hits(text)
                or_hits.extend(text_or_hits)
                or_hit_samples = collect_or_hit_samples(text, text_or_hits)
            else:
                or_hit_samples = {}
        else:
            or_hit_samples = {}
        return hits, "", sorted(set(or_hits), key=or_hits.index), or_hit_samples
    except Exception as err:
        return hits, str(err), sorted(set(or_hits), key=or_hits.index), {}


def build_result_rows(file_name, full_path, hits):
    rows = []
    for hit_type, hit in hits.items():
        rows.append({
            "risk": hit.get("risk", "未知"),
            "action": hit.get("action", "待確認"),
            "file_name": file_name,
            "hit_type": hit_type,
            "hit_count": hit.get("hit_count", 0),
            "sample": " / ".join(hit.get("samples", [])[:5]),
            "full_path": full_path,
        })
    return rows


def build_or_hit_rows(file_name, full_path, or_hits, existing_hit_types=None, samples_by_hit=None):
    existing_hit_types = set(existing_hit_types or [])
    samples_by_hit = samples_by_hit or {}
    rows = []
    for hit_type in or_hits:
        if hit_type in existing_hit_types:
            continue
        samples = samples_by_hit.get(hit_type) or ["OR運算命中"]
        rows.append({
            "risk": "提示",
            "action": "額外OR提示",
            "file_name": file_name,
            "hit_type": hit_type,
            "hit_count": len(samples),
            "sample": " / ".join(samples[:5]),
            "full_path": full_path,
        })
    return rows


def safe_zip_member_name(name):
    cleaned = name.replace("\\", "/").strip("/")
    parts = [part for part in cleaned.split("/") if part and part not in {".", ".."}]
    if not parts:
        return "zip-entry"
    return "/".join(parts)


def scan_zip_archive(path, display_path=None, depth=0, include_or_hits=False):
    archive_path = Path(path)
    archive_display_path = display_path or str(archive_path)
    rows = []
    or_hits = []
    total_uncompressed = 0

    if depth >= MAX_ZIP_SCAN_DEPTH:
        if include_or_hits:
            return rows, or_hits
        return rows

    try:
        with zipfile.ZipFile(archive_path, "r") as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue

                inner_name = safe_zip_member_name(info.filename)
                inner_display_path = f"{archive_display_path}::{inner_name}"
                inner_file_name = Path(inner_name).name

                if info.flag_bits & 0x1:
                    hits = {ENCRYPTED_FILE_RULE: make_policy_hit("高", 1, ["zip entry encrypted"])}
                    rows.extend(build_result_rows(inner_file_name, inner_display_path, hits))
                    or_hits.append(ENCRYPTED_FILE_RULE)
                    continue

                if info.file_size > MAX_ZIP_ENTRY_BYTES:
                    continue

                total_uncompressed += info.file_size
                if total_uncompressed > MAX_ZIP_TOTAL_BYTES:
                    break

                ext = Path(inner_name).suffix.lower()
                if ext not in SUPPORTED_SCAN_EXTS and ext != ".zip" and ext not in ENCRYPT_TARGET_EXTS:
                    continue

                try:
                    data = zf.read(info)
                except Exception:
                    continue

                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir) / inner_file_name
                    temp_path.write_bytes(data)

                    if ext == ".zip":
                        if include_or_hits:
                            nested_rows, nested_or_hits = scan_zip_archive(temp_path, inner_display_path, depth + 1, True)
                            rows.extend(nested_rows)
                            or_hits.extend(nested_or_hits)
                        else:
                            rows.extend(scan_zip_archive(temp_path, inner_display_path, depth + 1))
                    else:
                        if include_or_hits:
                            hits, _error, file_or_hits, file_or_hit_samples = scan_file_with_or_hits(temp_path)
                            or_hits.extend(file_or_hits)
                        else:
                            hits, _error = scan_file(temp_path)
                        rows.extend(build_result_rows(inner_file_name, inner_display_path, hits))
                        if include_or_hits:
                            rows.extend(build_or_hit_rows(
                                inner_file_name,
                                inner_display_path,
                                file_or_hits,
                                hits.keys(),
                                file_or_hit_samples,
                            ))
    except Exception:
        if include_or_hits:
            return rows, sorted(set(or_hits), key=or_hits.index)
        return rows

    if include_or_hits:
        return rows, sorted(set(or_hits), key=or_hits.index)
    return rows


def scan_path(path):
    path = Path(path)
    if path.suffix.lower() == ".zip":
        return scan_zip_archive(path)

    hits, _error = scan_file(path)
    return build_result_rows(path.name, str(path), hits)


def scan_path_with_or_hits(path):
    path = Path(path)
    if path.suffix.lower() == ".zip":
        return scan_zip_archive(path, include_or_hits=True)

    hits, _error, or_hits, or_hit_samples = scan_file_with_or_hits(path)
    rows = build_result_rows(path.name, str(path), hits)
    rows.extend(build_or_hit_rows(path.name, str(path), or_hits, hits.keys(), or_hit_samples))
    return rows, or_hits


def open_file_with_default_app(path):
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.call(["open", path])
        else:
            subprocess.call(["xdg-open", path])
    except Exception:
        pass


class DLPScannerApp(ttk.Frame):
    def __init__(self, master):
        super().__init__(master, padding=12)
        self.master = master
        self.pack(fill="both", expand=True)

        self.folder_var = StringVar()
        self.status_var = StringVar(value="請選擇要掃描的資料夾")
        self.include_subfolder_var = BooleanVar(value=True)
        self.auto_open_report_var = BooleanVar(value=False)

        self.message_queue = queue.Queue()
        self.scanning = False
        self.export_rows = []
        self.last_report_path = ""

        self.build_ui()

    def build_ui(self):
        self.master.title(f"{APP_NAME} v{APP_VERSION}")
        self.master.geometry("1120x680")
        self.master.minsize(960, 560)

        top_frame = ttk.Frame(self)
        top_frame.pack(fill="x", pady=(0, 8))

        ttk.Label(top_frame, text="掃描資料夾：").pack(side="left")
        ttk.Entry(top_frame, textvariable=self.folder_var).pack(side="left", fill="x", expand=True, padx=6)
        ttk.Button(top_frame, text="瀏覽", command=self.pick_folder).pack(side="left")

        self.scan_button = ttk.Button(top_frame, text="開始掃描", command=self.start_scan)
        self.scan_button.pack(side="left", padx=(6, 0))

        self.export_button = ttk.Button(top_frame, text="匯出 Excel", command=self.export_excel, state="disabled")
        self.export_button.pack(side="left", padx=(6, 0))

        option_frame = ttk.LabelFrame(self, text="掃描選項", padding=8)
        option_frame.pack(fill="x", pady=(0, 8))

        ttk.Checkbutton(option_frame, text="包含子資料夾", variable=self.include_subfolder_var).pack(side="left", padx=(0, 12))
        ttk.Checkbutton(option_frame, text="匯出後自動開啟 Excel", variable=self.auto_open_report_var).pack(side="left", padx=(0, 12))
        ttk.Label(option_frame, text="提示：信用卡號已加入 Luhn 檢核；畫面與報表會遮罩部分敏感內容。").pack(side="left")

        progress_frame = ttk.Frame(self)
        progress_frame.pack(fill="x", pady=(0, 8))
        self.progress = ttk.Progressbar(progress_frame, mode="determinate")
        self.progress.pack(side="left", fill="x", expand=True)
        ttk.Label(progress_frame, textvariable=self.status_var, width=45).pack(side="left", padx=8)

        result_frame = ttk.LabelFrame(self, text="掃描結果", padding=8)
        result_frame.pack(fill="both", expand=True)

        columns = ("risk", "action", "file_name", "hit_type", "hit_count", "sample", "full_path")
        headings = {
            "risk": "風險",
            "action": "處置",
            "file_name": "檔案名稱",
            "hit_type": "命中類型",
            "hit_count": "次數",
            "sample": "命中範例",
            "full_path": "完整路徑",
        }
        widths = {"risk": 70, "action": 110, "file_name": 220, "hit_type": 220, "hit_count": 60, "sample": 300, "full_path": 420}

        self.tree = ttk.Treeview(result_frame, columns=columns, show="headings", height=20)
        for col in columns:
            self.tree.heading(col, text=headings[col])
            self.tree.column(col, width=widths[col], anchor="w")

        y_scroll = ttk.Scrollbar(result_frame, orient="vertical", command=self.tree.yview)
        x_scroll = ttk.Scrollbar(result_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        y_scroll.grid(row=0, column=1, sticky="ns")
        x_scroll.grid(row=1, column=0, sticky="ew")
        result_frame.rowconfigure(0, weight=1)
        result_frame.columnconfigure(0, weight=1)

        self.tree.tag_configure("高", background="#ffd6d6")
        self.tree.tag_configure("中", background="#fff2cc")
        self.tree.tag_configure("低", background="#e2f0d9")
        self.tree.tag_configure("未分類", background="#eeeeee")

        bottom_frame = ttk.Frame(self)
        bottom_frame.pack(fill="x", pady=(8, 0))
        ttk.Label(bottom_frame, text="支援格式：txt, csv, log, md, json, xml, html, doc, docx, xls, xlsx, ppt, pdf, zip。").pack(side="left")

    def pick_folder(self):
        folder = filedialog.askdirectory(title="選擇要掃描的資料夾")
        if folder:
            self.folder_var.set(folder)
            self.status_var.set(f"已選擇：{folder}")

    def start_scan(self):
        if self.scanning:
            return

        folder = self.folder_var.get().strip()
        if not folder or not os.path.isdir(folder):
            messagebox.showwarning("提示", "請先選擇有效的資料夾。")
            return

        self.clear_results()
        self.scanning = True
        self.scan_button.config(state="disabled")
        self.export_button.config(state="disabled")

        worker = threading.Thread(
            target=self.scan_worker,
            args=(Path(folder), self.include_subfolder_var.get()),
            daemon=True,
        )
        worker.start()
        self.after(100, self.poll_queue)

    def clear_results(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.export_rows = []
        self.last_report_path = ""
        self.progress["value"] = 0
        self.status_var.set("準備掃描...")

    def scan_worker(self, folder, include_subfolder):
        files = [path for path in (folder.rglob("*") if include_subfolder else folder.glob("*")) if path.is_file()]
        total_files = len(files)
        self.message_queue.put({"type": "init", "total": total_files})

        hit_file_set = set()
        total_hit_rules = 0
        or_hit_names = []

        for index, file_path in enumerate(files, start=1):
            result_rows, file_or_hits = scan_path_with_or_hits(file_path)
            for hit_name in file_or_hits:
                if hit_name not in or_hit_names:
                    or_hit_names.append(hit_name)
            if result_rows:
                hit_file_set.update(row["full_path"] for row in result_rows)
                for row in result_rows:
                    total_hit_rules += 1
                    message = {"type": "row"}
                    message.update(row)
                    self.message_queue.put(message)

            self.message_queue.put({"type": "progress", "current": index, "total": total_files, "file_name": file_path.name})

        self.message_queue.put({
            "type": "done",
            "total": total_files,
            "hit_files": len(hit_file_set),
            "hit_rules": total_hit_rules,
            "or_hits": or_hit_names,
        })

    def poll_queue(self):
        try:
            while True:
                message = self.message_queue.get_nowait()
                message_type = message["type"]

                if message_type == "init":
                    self.progress.config(maximum=max(message["total"], 1), value=0)

                elif message_type == "progress":
                    self.progress["value"] = message["current"]
                    self.status_var.set(f"掃描中：{message['current']}/{message['total']} - {message['file_name']}")

                elif message_type == "row":
                    risk = message["risk"]
                    values = (
                        message["risk"], message["action"], message["file_name"], message["hit_type"],
                        message["hit_count"], message["sample"], message["full_path"],
                    )
                    self.tree.insert("", END, values=values, tags=(risk,))
                    self.export_rows.append({
                        "掃描時間": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "風險等級": message["risk"],
                        "處置": message["action"],
                        "檔案名稱": message["file_name"],
                        "命中類型": message["hit_type"],
                        "命中次數": message["hit_count"],
                        "命中範例": message["sample"],
                        "完整路徑": message["full_path"],
                    })

                elif message_type == "done":
                    self.scanning = False
                    self.scan_button.config(state="normal")
                    total = message["total"]
                    hit_files = message["hit_files"]
                    hit_rules = message["hit_rules"]
                    or_hits = message.get("or_hits", [])
                    self.status_var.set(f"完成：共掃描 {total} 個檔案，命中 {hit_files} 個檔案 / {hit_rules} 項規則")

                    if self.export_rows:
                        self.export_button.config(state="normal")

                    if or_hits:
                        messagebox.showwarning("OR運算命中", f"[OR運算命中:{', '.join(or_hits)}]")

                    if hit_files > 0:
                        messagebox.showwarning(
                            "命中提示",
                            f"掃描完成。\n\n共掃描 {total} 個檔案。\n"
                            f"發現 {hit_files} 個檔案命中敏感資料規則。\n"
                            f"合計 {hit_rules} 項規則命中。\n\n請查看下方清單或匯出 Excel 報表。",
                        )
                    elif not or_hits:
                        messagebox.showinfo("掃描完成", f"共掃描 {total} 個檔案，未發現命中。")
                    return
        except queue.Empty:
            pass

        if self.scanning:
            self.after(100, self.poll_queue)

    def export_excel(self):
        if not self.export_rows:
            messagebox.showwarning("提示", "目前沒有可匯出的掃描結果。")
            return

        if Workbook is None:
            messagebox.showerror("缺少套件", "尚未安裝 openpyxl，無法匯出 Excel。\n\n請執行：pip install openpyxl")
            return

        default_name = f"DLP掃描報表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        file_path = filedialog.asksaveasfilename(
            title="儲存 Excel 報表",
            defaultextension=".xlsx",
            initialfile=default_name,
            filetypes=[("Excel 檔案", "*.xlsx"), ("所有檔案", "*.*")],
        )
        if not file_path:
            return

        try:
            self.write_excel_report(file_path)
            self.last_report_path = file_path
            messagebox.showinfo("匯出完成", f"Excel 報表已匯出：\n{file_path}")
            if self.auto_open_report_var.get():
                open_file_with_default_app(file_path)
        except Exception as err:
            messagebox.showerror("匯出失敗", f"匯出 Excel 時發生錯誤：\n{err}")

    def write_excel_report(self, file_path):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "掃描結果"

        headers = ["掃描時間", "風險等級", "處置", "檔案名稱", "命中類型", "命中次數", "命中範例", "完整路徑"]
        sheet.append(headers)

        for row in self.export_rows:
            sheet.append([
                row.get("掃描時間", ""), row.get("風險等級", ""), row.get("處置", ""), row.get("檔案名稱", ""),
                row.get("命中類型", ""), row.get("命中次數", ""), row.get("命中範例", ""), row.get("完整路徑", ""),
            ])

        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        center = Alignment(horizontal="center", vertical="center")
        wrap = Alignment(wrap_text=True, vertical="top")

        for cell in sheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center

        risk_fill = {
            "高": PatternFill("solid", fgColor="FFC7CE"),
            "中": PatternFill("solid", fgColor="FFEB9C"),
            "低": PatternFill("solid", fgColor="C6EFCE"),
            "未分類": PatternFill("solid", fgColor="D9EAD3"),
        }

        for row_index in range(2, sheet.max_row + 1):
            risk = sheet.cell(row=row_index, column=2).value
            if risk in risk_fill:
                for col_index in range(1, sheet.max_column + 1):
                    sheet.cell(row=row_index, column=col_index).fill = risk_fill[risk]
            for col_index in range(1, sheet.max_column + 1):
                sheet.cell(row=row_index, column=col_index).alignment = wrap

        widths = {1: 20, 2: 10, 3: 14, 4: 28, 5: 30, 6: 10, 7: 55, 8: 80}
        for col_index, width in widths.items():
            sheet.column_dimensions[get_column_letter(col_index)].width = width

        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions

        summary = workbook.create_sheet("摘要")
        summary["A1"] = "DLP 掃描報表摘要"
        summary["A1"].font = Font(size=16, bold=True)
        summary_data = [
            ("產生時間", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("命中紀錄數", len(self.export_rows)),
            ("高風險筆數", sum(1 for row in self.export_rows if row.get("風險等級") == "高")),
            ("中風險筆數", sum(1 for row in self.export_rows if row.get("風險等級") == "中")),
            ("低風險筆數", sum(1 for row in self.export_rows if row.get("風險等級") == "低")),
        ]
        for index, item in enumerate(summary_data, start=3):
            summary.cell(row=index, column=1).value = item[0]
            summary.cell(row=index, column=2).value = item[1]
        summary.column_dimensions["A"].width = 20
        summary.column_dimensions["B"].width = 30

        workbook.save(file_path)


def main():
    root = Tk()
    try:
        style = ttk.Style()
        available_themes = style.theme_names()
        if "vista" in available_themes:
            style.theme_use("vista")
        elif "clam" in available_themes:
            style.theme_use("clam")
    except Exception:
        pass

    DLPScannerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
